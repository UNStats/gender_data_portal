from os import listdir
from os.path import isfile, join
from docx import Document
import utils

# List of availability_ts files
datafiles = [f for f in listdir('metadata_temp/') if isfile(join('metadata_temp/', f))]


for fdx, f in enumerate(datafiles):

    # if fdx != 0:
    #     continue

    x = f.split('__')

    indicator_label = x[0].replace('_', '.')
    indicator_id = x[1]
    if x[2] == 'metadata.docx':
        series_id = None
    else: 
        series_id = x[2]

    print(f"{fdx}")
    print(f"{indicator_label=}")
    print(f"{indicator_id=}")
    print(f"{series_id=}")
    print("----")


    wordDoc = Document(join('metadata_temp/', f))

    tables = wordDoc.tables

    md = ''
    for tdx, table in enumerate(wordDoc.tables):

        print(f"Table = {tdx}")


        for rdx, row in enumerate(table.rows):
            if rdx == 0:
                continue

            md = md + '## ' + utils.clean_str(row.cells[2].text) + '\n\n'
            md = md + utils.clean_str(row.cells[3].text)+ '\n\n' + '---'+ '\n\n'



    #open text file
    text_file = open("temp/" + f.replace('.docx', '.md'), "w")
    text_file.write(md)
    text_file.close()